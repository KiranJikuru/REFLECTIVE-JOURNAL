import os
import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
import re
import pypandoc
from dotenv import load_dotenv
from docx2pdf import convert

load_dotenv()

TEMPLATE_PATH = "vam.docx"

PREFERRED_MODELS = [
    "gemini-2.0-flash",
    "gemini-2.0-pro",
    "gemini-1.5-flash",
    "gemini-1.5-pro"
]

WORDS = {
    "experiences": 150,
    "feelings": 150,
    "insights": 300,
    "conclusion": 100,
    "applications_count": 5
}


# -------------------------------
# Initialize Gemini
# -------------------------------
def init_genai():
    key = os.environ.get("GOOGLE_API_KEY")
    if not key:
        st.error("GOOGLE_API_KEY not found")
        st.stop()

    genai.configure(api_key=key)

    for m in PREFERRED_MODELS:
        try:
            model = genai.GenerativeModel(m)
            model.generate_content("ping")
            return m
        except:
            continue

    st.error("No working Gemini model available.")
    st.stop()


# -------------------------------
# Gemini Call
# -------------------------------
def call_gemini(prompt, model, max_tokens=900):
    obj = genai.GenerativeModel(model)
    res = obj.generate_content(prompt, generation_config={"max_output_tokens": max_tokens})
    return res.text.strip()


def enforce_count(txt, count):
    txt = re.sub(r"\s+", " ", txt)
    words = txt.split()
    if len(words) <= count:
        return txt
    return " ".join(words[:count]) + "..."


# -------------------------------
# Generate Reflective Sections
# -------------------------------
def generate_section(model, title, topic, nwords):
    prompt = f"""
Write a simple and easy-to-understand reflective paragraph of about {nwords} words.

The paragraph MUST start exactly like this:
"In this module, I have learned that {topic}..."

Writing rules:
• Use simple English only.
• No complex or academic words.
• A single clear paragraph.
• Make it personal and human-like.
• Describe real feelings about what was learned.
• No bullet points.
"""
    raw = call_gemini(prompt, model)
    return enforce_count(raw, nwords)


def generate_apps(model, topic, n=5):
    prompt = f"""
Give {n} simple real-life applications of {topic}.
Rules:
• One short sentence each.
• Use simple English.
• No bullets, no symbols.
• No asterisks (*).
• Output as plain lines separated by newlines.
"""
    raw = call_gemini(prompt, model)

    lines = [re.sub(r"^[0-9\.\-\*\•\s]+", "", l).strip() for l in raw.split("\n") if l.strip()]

    while len(lines) < n:
        lines.append(f"A simple example of {topic}.")

    clean_lines = lines[:n]

    numbered = []
    for i, line in enumerate(clean_lines, start=1):
        numbered.append(f"{i}. {line}")

    return "\n".join(numbered)


# -------------------------------
# Apply Formatting (TNR, 12pt, justify, 1.5 spacing)
# -------------------------------
def apply_font_styles(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(14)

        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.7

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        run.font.size = Pt(12)

                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph.paragraph_format.line_spacing = 1.5


# -------------------------------
# Fill DOCX Template
# -------------------------------
def fill_template(data):
    if not os.path.exists(TEMPLATE_PATH):
        st.error(f"Template not found: {TEMPLATE_PATH}")
        st.stop()

    doc = Document(TEMPLATE_PATH)

    # Replace placeholders
    for p in doc.paragraphs:
        for key, value in data.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    apply_font_styles(doc)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# -------------------------------
# Convert DOCX to PDF
# -------------------------------

import tempfile
import os

def convert_to_pdf(docx_bytes):
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "temp.docx")
        pdf_path = os.path.join(tmpdir, "temp.pdf")

        # Save DOCX temporarily
        with open(docx_path, "wb") as f:
            f.write(docx_bytes.getvalue())

        # Convert to PDF
        convert(docx_path, pdf_path)

        # Read the PDF back
        with open(pdf_path, "rb") as f:
            return f.read()



# -------------------------------
# Streamlit UI
# -------------------------------
st.title("Reflective Journal Generator")

assignment_name = st.text_input("Assignment Name")
student_name = st.text_input("Student Name")
rollno = st.text_input("Roll Number")
class_section = st.text_input("Class / Section")
level = st.text_input("Level")
year_term = st.text_input("Year / Term")
journal_title = st.text_input("Journal Entry Title")
topic = st.text_input("Enter Topic to Generate")
subject_name = st.text_input("Subject Name")


if st.button("Generate RJ"):
    if not topic:
        st.error("Please enter a topic.")
        st.stop()

    model = init_genai()
    st.success(f"Using model: {model}")

    exp = generate_section(model, "Experiences", topic, WORDS["experiences"])
    feel = generate_section(model, "Feelings", topic, WORDS["feelings"])
    ins = generate_section(model, "Insights", topic, WORDS["insights"])
    conc = generate_section(model, "Conclusion", topic, WORDS["conclusion"])
    apps = generate_apps(model, topic)

    final = {
        "{{assignment_name}}": assignment_name,
        "{{student_name}}": student_name,
        "{{rollno}}": rollno,
        "{{class_section}}": class_section,
        "{{level}}": level,
        "{{year_term}}": year_term,
        "{{subject_name}}": subject_name,
        "{{title}}": journal_title,
        "{{experiences}}": exp,
        "{{feelings}}": feel,
        "{{insights}}": ins,
        "{{applications}}": apps,
        "{{conclusion}}": conc,
    }

    docx_output = fill_template(final)

    st.download_button("Download DOCX", data=docx_output, file_name="Journal_Output.docx")

    # Convert to PDF
    pdf_bytes = convert_to_pdf(docx_output)
    st.download_button("Download PDF", data=pdf_bytes, file_name="Journal_Output.pdf")
